import fitz  # PyMuPDF
import docx
from docx.shared import Inches
import os
import tempfile

# Open PDF
pdf_path = "geschiedenis.pdf"
output_docx = "extrait_du_pdf_geschiedenis.docx"

if not os.path.exists(pdf_path):
    raise FileNotFoundError(f"Le fichier PDF '{pdf_path}' est introuvable.")

doc = docx.Document()
pdf = fitz.open(pdf_path)

# Create a temporary folder to store extracted images
with tempfile.TemporaryDirectory() as tmpdir:
    for page_num, page in enumerate(pdf):
        # Texte
        text = page.get_text("text")
        doc.add_heading(f"Page {page_num + 1}", level=1)

        if text.strip():
            doc.add_paragraph(text)
        else:
            doc.add_paragraph("(Aucun texte détecté sur cette page)")

        # Images
        images = page.get_images(full=True)
        if images:
            doc.add_paragraph("Images extraites :")
            for img_index, img in enumerate(images):
                try:
                    xref = img[0]
                    base_image = pdf.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]

                    # Skip unsupported formats for Word (like .jp2)
                    if image_ext.lower() not in ["png", "jpeg", "jpg"]:
                        continue

                    image_filename = os.path.join(
                        tmpdir, f"image_{page_num+1}_{img_index+1}.{image_ext}"
                    )
                    with open(image_filename, "wb") as f:
                        f.write(image_bytes)

                    doc.add_picture(image_filename, width=Inches(5.0))
                except Exception as e:
                    doc.add_paragraph(f"[Erreur lors de l'insertion d'une image : {e}]")
        else:
            doc.add_paragraph("Aucune image sur cette page.")

# Sauvegarde finale
doc.save("output.docx")
